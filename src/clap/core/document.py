"""Document loading utilities for CLAP."""

import base64
import re
from dataclasses import dataclass


@dataclass
class Document:
    """Simple document class."""

    content: str
    metadata: dict


def load_document(file_path: str) -> list[Document]:
    """Load document based on file type."""
    lower = file_path.lower()
    if lower.endswith(".pdf"):
        return _load_pdf(file_path)
    elif lower.endswith((".doc", ".docx")):
        return _load_docx(file_path)
    elif lower.endswith((".md", ".markdown")):
        return _load_markdown(file_path)
    elif lower.endswith(".txt"):
        return _load_text(file_path)
    elif lower.endswith(".svg"):
        return _load_svg(file_path)
    elif lower.endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp")):
        return _load_image(file_path)
    return []


def _load_pdf(file_path: str) -> list[Document]:
    """Load PDF file using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError:
        print("pypdf not installed: pip install pypdf")
        return []

    reader = PdfReader(file_path)
    docs = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            docs.append(Document(content=text, metadata={"source": file_path, "page": i}))
    return docs


def _load_docx(file_path: str) -> list[Document]:
    """Load DOCX file using python-docx."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        print("python-docx not installed: pip install python-docx")
        return []

    doc = DocxDocument(file_path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text)

    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text:
                tables_text.append(row_text)

    if tables_text:
        text += "\n\n" + "\n".join(tables_text)

    return [Document(content=text, metadata={"source": file_path, "type": "docx"})] if text else []


def _load_markdown(file_path: str) -> list[Document]:
    """Load Markdown file."""
    with open(file_path, encoding="utf-8") as f:
        text = f.read()
    return [Document(content=text, metadata={"source": file_path, "type": "markdown"})] if text else []


def _load_svg(file_path: str) -> list[Document]:
    """Load SVG file, extract text content and prepare for vision models."""
    with open(file_path, encoding="utf-8") as f:
        svg_content = f.read()

    text_elements = re.findall(r"<text[^>]*>([^<]*)</text>", svg_content)
    title_match = re.search(r"<title>([^<]*)</title>", svg_content)
    desc_match = re.search(r"<desc>([^<]*)</desc>", svg_content)

    extracted_text = []
    if title_match:
        extracted_text.append(f"Title: {title_match.group(1)}")
    if desc_match:
        extracted_text.append(f"Description: {desc_match.group(1)}")
    if text_elements:
        cleaned = [t.strip() for t in text_elements if t.strip()]
        if cleaned:
            extracted_text.append("Text content: " + " | ".join(cleaned))

    svg_base64 = base64.b64encode(svg_content.encode("utf-8")).decode("utf-8")

    return [
        Document(
            content="\n".join(extracted_text) if extracted_text else "",
            metadata={
                "source": file_path,
                "type": "svg",
                "svg_base64": svg_base64,
                "raw_svg": svg_content,
            },
        )
    ]


def _load_text(file_path: str) -> list[Document]:
    """Load plain text file."""
    with open(file_path, encoding="utf-8") as f:
        text = f.read()
    return [Document(content=text, metadata={"source": file_path})] if text else []


def _load_image(file_path: str) -> list[Document]:
    """Load image for vision models."""
    try:
        from PIL import Image

        with Image.open(file_path) as img:
            img.verify()
        return [Document(content="", metadata={"source": file_path, "type": "image"})]
    except Exception:
        return []


def is_image(file_path: str) -> bool:
    """Check if file is an image."""
    return file_path.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"))


def is_svg(file_path: str) -> bool:
    """Check if file is an SVG."""
    return file_path.lower().endswith(".svg")


def chunk_text(text: str, chunk_size: int = 2000, overlap: int = 200) -> list[str]:
    """Split text into overlapping chunks."""
    if len(text) <= chunk_size:
        return [text] if text else []

    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start = end - overlap
    return chunks
